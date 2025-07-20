#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from flask import Flask, request, jsonify, send_from_directory, render_template_string, send_file
from flask_cors import CORS
from flask_mail import Mail, Message
from datetime import datetime, timedelta
import json
import os
import uuid
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import Workbook
from openpyxl.styles import Font, Fill, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import io
import base64

# Configura√ß√£o do Flask
app = Flask(__name__)
app.config['JSON_AS_ASCII'] = False
app.config['JSONIFY_MIMETYPE'] = 'application/json; charset=utf-8'
CORS(app)

# Configura√ß√£o de Email (compat√≠vel com suas vari√°veis)
app.config['MAIL_SERVER'] = os.environ.get('EMAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT'] = int(os.environ.get('EMAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.environ.get('EMAIL_USER', os.environ.get('MAIL_USERNAME', ''))
app.config['MAIL_PASSWORD'] = os.environ.get('EMAIL_PASS', os.environ.get('MAIL_PASSWORD', ''))
app.config['MAIL_DEFAULT_SENDER'] = os.environ.get('EMAIL_USER', os.environ.get('MAIL_USERNAME', ''))

mail = Mail(app)

# Configura√ß√£o de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Diret√≥rios
PROPOSTAS_DIR = 'propostas'
STATIC_DIR = 'static'

# Criar diret√≥rios se n√£o existirem
for dir_name in [PROPOSTAS_DIR, STATIC_DIR]:
    try:
        if not os.path.exists(dir_name):
            os.makedirs(dir_name)
            logger.info(f"Diret√≥rio {dir_name} criado")
    except Exception as e:
        logger.error(f"Erro ao criar diret√≥rio {dir_name}: {e}")

# Base de dados em mem√≥ria
propostas_db = {}
processos_db = {}
fornecedores_db = {}

# Configura√ß√µes de email
EMAIL_CONFIG = {
    'destinatario_principal': os.environ.get('EMAIL_SUPRIMENTOS', 'portaldofornecedor.arias@gmail.com')
}

def inicializar_dados_exemplo():
    """Inicializa dados de exemplo para demonstra√ß√£o"""
    global processos_db, fornecedores_db
    
    # Dados de exemplo de processos
    agora = datetime.now()
    processos_exemplo = [
        {
            "numero": "001/2025",
            "objeto": "Constru√ß√£o de Escola Municipal",
            "modalidade": "Concorr√™ncia",
            "prazo": (agora + timedelta(days=15)).isoformat(),
            "valor_estimado": "R$ 2.500.000,00",
            "status": "ativo"
        },
        {
            "numero": "002/2025", 
            "objeto": "Reforma do Centro de Sa√∫de",
            "modalidade": "Tomada de Pre√ßos",
            "prazo": (agora + timedelta(days=10)).isoformat(),
            "valor_estimado": "R$ 850.000,00",
            "status": "ativo"
        },
        {
            "numero": "003/2025",
            "objeto": "Pavimenta√ß√£o de Ruas",
            "modalidade": "Concorr√™ncia", 
            "prazo": (agora + timedelta(days=20)).isoformat(),
            "valor_estimado": "R$ 1.200.000,00",
            "status": "ativo"
        }
    ]
    
    # Adicionar aos processos_db
    for processo in processos_exemplo:
        processos_db[processo["numero"]] = processo
    
    # Dados de exemplo de fornecedores
    fornecedores_exemplo = [
        {
            "cnpj": "12.345.678/0001-90",
            "razaoSocial": "Construtora Exemplo LTDA",
            "email": "contato@exemplo.com.br",
            "telefone": "(11) 99999-9999",
            "status": "ativo"
        }
    ]
    
    # Adicionar aos fornecedores_db
    for fornecedor in fornecedores_exemplo:
        fornecedores_db[fornecedor["cnpj"]] = fornecedor
    
    logger.info(f"Dados de exemplo inicializados: {len(processos_db)} processos, {len(fornecedores_db)} fornecedores")

# Inicializar dados de exemplo
inicializar_dados_exemplo()

def gerar_excel_proposta(dados_proposta):
    """Gera arquivo Excel com a proposta comercial"""
    wb = Workbook()
    
    # Estilos
    header_font = Font(name='Arial', size=12, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    
    title_font = Font(name='Arial', size=14, bold=True)
    subtitle_font = Font(name='Arial', size=11, bold=True)
    normal_font = Font(name='Arial', size=10)
    
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Aba 1: Resumo
    ws_resumo = wb.active
    ws_resumo.title = "Resumo"
    
    # Cabe√ßalho
    ws_resumo.merge_cells('A1:F1')
    ws_resumo['A1'] = 'PROPOSTA COMERCIAL'
    ws_resumo['A1'].font = Font(name='Arial', size=16, bold=True)
    ws_resumo['A1'].alignment = Alignment(horizontal="center")
    
    ws_resumo['A3'] = 'Dados da Empresa'
    ws_resumo['A3'].font = subtitle_font
    
    # Dados da empresa
    dados_empresa = [
        ['Raz√£o Social:', dados_proposta.get('dados', {}).get('razaoSocial', '')],
        ['CNPJ:', dados_proposta.get('dados', {}).get('cnpj', '')],
        ['Endere√ßo:', dados_proposta.get('dados', {}).get('endereco', '')],
        ['Cidade:', dados_proposta.get('dados', {}).get('cidade', '')],
        ['Telefone:', dados_proposta.get('dados', {}).get('telefone', '')],
        ['E-mail:', dados_proposta.get('dados', {}).get('email', '')],
        ['Respons√°vel T√©cnico:', dados_proposta.get('dados', {}).get('respTecnico', '')],
        ['CREA/CAU:', dados_proposta.get('dados', {}).get('crea', '')]
    ]
    
    row = 5
    for label, value in dados_empresa:
        ws_resumo[f'A{row}'] = label
        ws_resumo[f'A{row}'].font = Font(bold=True)
        ws_resumo[f'B{row}'] = value
        ws_resumo.merge_cells(f'B{row}:D{row}')
        row += 1
    
    # Resumo financeiro
    ws_resumo[f'A{row+2}'] = 'Resumo Financeiro'
    ws_resumo[f'A{row+2}'].font = subtitle_font
    
    row += 4
    resumo_financeiro = [
        ['M√£o de Obra:', dados_proposta.get('comercial', {}).get('totalMaoObra', '0,00')],
        ['Materiais:', dados_proposta.get('comercial', {}).get('totalMateriais', '0,00')],
        ['Equipamentos:', dados_proposta.get('comercial', {}).get('totalEquipamentos', '0,00')],
        ['BDI:', f"{dados_proposta.get('comercial', {}).get('bdiPercentual', '0')}% = R$ {dados_proposta.get('comercial', {}).get('bdiValor', '0,00')}"],
        ['VALOR TOTAL:', f"R$ {dados_proposta.get('comercial', {}).get('valorTotal', '0,00')}"]
    ]
    
    for label, value in resumo_financeiro:
        ws_resumo[f'A{row}'] = label
        ws_resumo[f'A{row}'].font = Font(bold=True)
        ws_resumo[f'B{row}'] = value
        ws_resumo.merge_cells(f'B{row}:D{row}')
        if label == 'VALOR TOTAL:':
            ws_resumo[f'A{row}'].font = Font(bold=True, size=12, color="FF0000")
            ws_resumo[f'B{row}'].font = Font(bold=True, size=12, color="FF0000")
        row += 1
    
    # Aba 2: Servi√ßos
    if dados_proposta.get('comercial', {}).get('servicos'):
        ws_servicos = wb.create_sheet("Servi√ßos")
        
        # Cabe√ßalho
        headers = ['Item', 'Descri√ß√£o', 'Unidade', 'Quantidade', 'Pre√ßo Unit.', 'Total']
        for col, header in enumerate(headers, 1):
            cell = ws_servicos.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Dados
        row = 2
        for servico in dados_proposta['comercial']['servicos']:
            for col, value in enumerate(servico[:6], 1):
                cell = ws_servicos.cell(row=row, column=col, value=value)
                cell.border = border
                if col >= 4:  # Colunas num√©ricas
                    cell.alignment = Alignment(horizontal="right")
            row += 1
        
        # Total
        ws_servicos[f'E{row}'] = 'TOTAL:'
        ws_servicos[f'E{row}'].font = Font(bold=True)
        ws_servicos[f'F{row}'] = f"R$ {dados_proposta.get('comercial', {}).get('totalServicos', '0,00')}"
        ws_servicos[f'F{row}'].font = Font(bold=True)
        
        # Ajustar largura das colunas
        ws_servicos.column_dimensions['B'].width = 50
        for col in ['C', 'D', 'E', 'F']:
            ws_servicos.column_dimensions[col].width = 15
    
    # Aba 3: M√£o de Obra
    if dados_proposta.get('comercial', {}).get('maoObra'):
        ws_mo = wb.create_sheet("M√£o de Obra")
        
        headers = ['Fun√ß√£o', 'Qtd', 'Tempo', 'Sal√°rio', 'Encargos %', 'Total Mensal', 'Total Geral']
        for col, header in enumerate(headers, 1):
            cell = ws_mo.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        row = 2
        for mo in dados_proposta['comercial']['maoObra']:
            for col, value in enumerate(mo[:7], 1):
                cell = ws_mo.cell(row=row, column=col, value=value)
                cell.border = border
                if col >= 2:
                    cell.alignment = Alignment(horizontal="right")
            row += 1
        
        ws_mo[f'F{row}'] = 'TOTAL:'
        ws_mo[f'F{row}'].font = Font(bold=True)
        ws_mo[f'G{row}'] = f"R$ {dados_proposta.get('comercial', {}).get('totalMaoObra', '0,00')}"
        ws_mo[f'G{row}'].font = Font(bold=True)
    
    # Aba 4: Materiais
    if dados_proposta.get('comercial', {}).get('materiaisComercial'):
        ws_mat = wb.create_sheet("Materiais")
        
        headers = ['Material', 'Especifica√ß√£o', 'Unidade', 'Quantidade', 'Pre√ßo Unit.', 'Total']
        for col, header in enumerate(headers, 1):
            cell = ws_mat.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        row = 2
        for material in dados_proposta['comercial']['materiaisComercial']:
            for col, value in enumerate(material[:6], 1):
                cell = ws_mat.cell(row=row, column=col, value=value)
                cell.border = border
                if col >= 4:
                    cell.alignment = Alignment(horizontal="right")
            row += 1
        
        ws_mat[f'E{row}'] = 'TOTAL:'
        ws_mat[f'E{row}'].font = Font(bold=True)
        ws_mat[f'F{row}'] = f"R$ {dados_proposta.get('comercial', {}).get('totalMateriais', '0,00')}"
        ws_mat[f'F{row}'].font = Font(bold=True)
    
    # Aba 5: Equipamentos
    if dados_proposta.get('comercial', {}).get('equipamentosComercial'):
        ws_equip = wb.create_sheet("Equipamentos")
        
        headers = ['Equipamento', 'Especifica√ß√£o', 'Quantidade', 'Tempo', 'Pre√ßo/m√™s', 'Total']
        for col, header in enumerate(headers, 1):
            cell = ws_equip.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        row = 2
        for equip in dados_proposta['comercial']['equipamentosComercial']:
            for col, value in enumerate(equip[:6], 1):
                cell = ws_equip.cell(row=row, column=col, value=value)
                cell.border = border
                if col >= 3:
                    cell.alignment = Alignment(horizontal="right")
            row += 1
        
        ws_equip[f'E{row}'] = 'TOTAL:'
        ws_equip[f'E{row}'].font = Font(bold=True)
        ws_equip[f'F{row}'] = f"R$ {dados_proposta.get('comercial', {}).get('totalEquipamentos', '0,00')}"
        ws_equip[f'F{row}'].font = Font(bold=True)
    
    # Aba 6: BDI
    if dados_proposta.get('comercial', {}).get('bdi'):
        ws_bdi = wb.create_sheet("BDI")
        
        headers = ['Componente', 'Percentual %', 'Valor R$']
        for col, header in enumerate(headers, 1):
            cell = ws_bdi.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        row = 2
        for bdi_item in dados_proposta['comercial']['bdi']:
            for col, value in enumerate(bdi_item[:3], 1):
                cell = ws_bdi.cell(row=row, column=col, value=value)
                cell.border = border
                if col >= 2:
                    cell.alignment = Alignment(horizontal="right")
            row += 1
        
        # Total BDI
        ws_bdi[f'A{row+1}'] = 'BDI TOTAL:'
        ws_bdi[f'B{row+1}'] = f"{dados_proposta.get('comercial', {}).get('bdiPercentual', '0')}%"
        ws_bdi[f'C{row+1}'] = f"R$ {dados_proposta.get('comercial', {}).get('bdiValor', '0,00')}"
        for col in range(1, 4):
            ws_bdi.cell(row=row+1, column=col).font = Font(bold=True, size=12)
    
    # Salvar em mem√≥ria
    excel_buffer = io.BytesIO()
    wb.save(excel_buffer)
    excel_buffer.seek(0)
    
    return excel_buffer

def gerar_word_proposta(dados_proposta):
    """Gera arquivo Word com a proposta t√©cnica"""
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # T√≠tulo principal
    titulo = doc.add_heading('PROPOSTA T√âCNICA', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dados da empresa
    doc.add_heading('1. DADOS DA EMPRESA', level=1)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    dados_empresa = [
        ('Raz√£o Social:', dados_proposta.get('dados', {}).get('razaoSocial', '')),
        ('CNPJ:', dados_proposta.get('dados', {}).get('cnpj', '')),
        ('Endere√ßo:', dados_proposta.get('dados', {}).get('endereco', '')),
        ('Cidade/UF:', dados_proposta.get('dados', {}).get('cidade', '')),
        ('Telefone:', dados_proposta.get('dados', {}).get('telefone', '')),
        ('E-mail:', dados_proposta.get('dados', {}).get('email', '')),
        ('Respons√°vel T√©cnico:', dados_proposta.get('dados', {}).get('respTecnico', '')),
        ('CREA/CAU:', dados_proposta.get('dados', {}).get('crea', ''))
    ]
    
    for i, (label, value) in enumerate(dados_empresa):
        table.cell(i, 0).text = label
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = str(value)
    
    # Objeto
    doc.add_heading('2. OBJETO DA CONCORR√äNCIA', level=1)
    doc.add_paragraph(dados_proposta.get('tecnica', {}).get('objetoConcorrencia', ''))
    
    # Escopo
    if dados_proposta.get('tecnica', {}).get('escopoInclusos'):
        doc.add_heading('3. ESCOPO DOS SERVI√áOS', level=1)
        doc.add_heading('3.1 Servi√ßos Inclusos:', level=2)
        doc.add_paragraph(dados_proposta['tecnica']['escopoInclusos'])
    
    if dados_proposta.get('tecnica', {}).get('escopoExclusos'):
        doc.add_heading('3.2 Servi√ßos Exclusos:', level=2)
        doc.add_paragraph(dados_proposta['tecnica']['escopoExclusos'])
    
    # Metodologia
    if dados_proposta.get('tecnica', {}).get('metodologia'):
        doc.add_heading('4. METODOLOGIA DE EXECU√á√ÉO', level=1)
        doc.add_paragraph(dados_proposta['tecnica']['metodologia'])
    
    if dados_proposta.get('tecnica', {}).get('sequenciaExecucao'):
        doc.add_heading('4.1 Sequ√™ncia de Execu√ß√£o:', level=2)
        doc.add_paragraph(dados_proposta['tecnica']['sequenciaExecucao'])
    
    # Prazo
    doc.add_heading('5. PRAZO DE EXECU√á√ÉO', level=1)
    prazo_exec = dados_proposta.get('tecnica', {}).get('prazoExecucao', 'N√£o informado')
    doc.add_paragraph(f'Prazo total para execu√ß√£o dos servi√ßos: {prazo_exec}')
    
    if dados_proposta.get('tecnica', {}).get('prazoMobilizacao'):
        doc.add_paragraph(f'Prazo de mobiliza√ß√£o: {dados_proposta["tecnica"]["prazoMobilizacao"]}')
    
    # Cronograma
    if dados_proposta.get('tecnica', {}).get('cronograma'):
        doc.add_heading('6. CRONOGRAMA DE EXECU√á√ÉO', level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Cabe√ßalho
        hdr_cells = table.rows[0].cells
        headers = ['Atividade', 'Dura√ß√£o', 'In√≠cio', 'Fim']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Dados
        for atividade in dados_proposta['tecnica']['cronograma']:
            row_cells = table.add_row().cells
            for i in range(min(4, len(atividade))):
                row_cells[i].text = str(atividade[i])
    
    # Equipe t√©cnica
    if dados_proposta.get('tecnica', {}).get('equipe'):
        doc.add_heading('7. EQUIPE T√âCNICA', level=1)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['Fun√ß√£o', 'Quantidade', 'Tempo (meses)']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        for membro in dados_proposta['tecnica']['equipe']:
            row_cells = table.add_row().cells
            for i in range(min(3, len(membro))):
                row_cells[i].text = str(membro[i])
    
    # Experi√™ncia
    if dados_proposta.get('tecnica', {}).get('experiencia'):
        doc.add_heading('8. EXPERI√äNCIA DA EMPRESA', level=1)
        
        for i, obra in enumerate(dados_proposta['tecnica']['experiencia'], 1):
            doc.add_heading(f'8.{i} Obra {i}:', level=2)
            if len(obra) > 0 and obra[0]:
                doc.add_paragraph(f'Obra: {obra[0]}')
            if len(obra) > 1 and obra[1]:
                doc.add_paragraph(f'Cliente: {obra[1]}')
            if len(obra) > 2 and obra[2]:
                doc.add_paragraph(f'Valor: {obra[2]}')
            if len(obra) > 3 and obra[3]:
                doc.add_paragraph(f'Ano: {obra[3]}')
    
    # Garantias
    if dados_proposta.get('tecnica', {}).get('garantias'):
        doc.add_heading('9. GARANTIAS OFERECIDAS', level=1)
        doc.add_paragraph(dados_proposta['tecnica']['garantias'])
    
    # Rodap√© com data
    doc.add_paragraph()
    doc.add_paragraph()
    data_proposta = doc.add_paragraph(f'Data da Proposta: {datetime.now().strftime("%d/%m/%Y")}')
    data_proposta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar em mem√≥ria
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    
    return word_buffer

def enviar_email_proposta(protocolo, dados_proposta):
    """Envia email com os dados da proposta e anexos Excel e Word"""
    if not app.config['MAIL_USERNAME'] or not app.config['MAIL_PASSWORD']:
        logger.warning("Configura√ß√µes de email n√£o definidas")
        return False
    
    try:
        # Gerar anexos
        excel_buffer = gerar_excel_proposta(dados_proposta)
        word_buffer = gerar_word_proposta(dados_proposta)
        
        # Criar mensagem
        msg = Message(
            subject=f"Nova Proposta Recebida - {protocolo}",
            recipients=[EMAIL_CONFIG['destinatario_principal']]
        )
        
        # Extrai informa√ß√µes principais
        empresa = dados_proposta.get('dados', {}).get('razaoSocial', 'N/A')
        cnpj = dados_proposta.get('dados', {}).get('cnpj', 'N/A')
        valor_total = dados_proposta.get('comercial', {}).get('valorTotal', '0,00')
        processo = dados_proposta.get('processo', 'N/A')
        prazo = dados_proposta.get('tecnica', {}).get('prazoExecucao', 'N/A')
        
        # Corpo do email em HTML
        msg.html = f"""
        <html>
        <body style="font-family: Arial, sans-serif;">
            <div style="background-color: #f0f0f0; padding: 20px; border-radius: 10px;">
                <h2 style="color: #333;">üè¢ Nova Proposta Recebida</h2>
                
                <div style="background-color: white; padding: 20px; border-radius: 5px; margin-top: 20px;">
                    <h3 style="color: #2c3e50;">Informa√ß√µes Principais</h3>
                    
                    <table style="width: 100%; border-collapse: collapse;">
                        <tr>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;"><strong>Protocolo:</strong></td>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;">{protocolo}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;"><strong>Processo:</strong></td>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;">{processo}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;"><strong>Empresa:</strong></td>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;">{empresa}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;"><strong>CNPJ:</strong></td>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;">{cnpj}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;"><strong>E-mail:</strong></td>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;">{dados_proposta.get('dados', {}).get('email', 'N/A')}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;"><strong>Telefone:</strong></td>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;">{dados_proposta.get('dados', {}).get('telefone', 'N/A')}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;"><strong>Prazo de Execu√ß√£o:</strong></td>
                            <td style="padding: 8px; border-bottom: 1px solid #ddd;">{prazo}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px;"><strong>Valor Total:</strong></td>
                            <td style="padding: 8px; color: #27ae60; font-size: 18px;"><strong>R$ {valor_total}</strong></td>
                        </tr>
                    </table>
                </div>
                
                <div style="margin-top: 20px; padding: 15px; background-color: #e8f5e9; border-radius: 5px;">
                    <p style="margin: 0;"><strong>üìé Anexos:</strong></p>
                    <ul>
                        <li>Proposta Comercial Detalhada (Excel)</li>
                        <li>Proposta T√©cnica Completa (Word)</li>
                    </ul>
                </div>
                
                <p style="text-align: center; color: #666; margin-top: 20px;">
                    <small>Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}</small>
                </p>
            </div>
        </body>
        </html>
        """
        
        # Anexar Excel
        msg.attach(
            f"proposta_comercial_{protocolo}.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            excel_buffer.getvalue()
        )
        
        # Anexar Word
        msg.attach(
            f"proposta_tecnica_{protocolo}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            word_buffer.getvalue()
        )
        
        # Enviar
        mail.send(msg)
        
        logger.info(f"Email enviado com sucesso para proposta {protocolo}")
        
        # Enviar email de confirma√ß√£o ao fornecedor
        if dados_proposta.get('dados', {}).get('email'):
            enviar_email_confirmacao_fornecedor(protocolo, dados_proposta)
        
        return True
        
    except Exception as e:
        logger.error(f"Erro ao enviar email: {str(e)}")
        return False

def enviar_email_confirmacao_fornecedor(protocolo, dados_proposta):
    """Envia email de confirma√ß√£o para o fornecedor"""
    try:
        email_fornecedor = dados_proposta.get('dados', {}).get('email')
        if not email_fornecedor:
            return False
        
        msg = Message(
            subject=f"Confirma√ß√£o de Recebimento - Proposta {protocolo}",
            recipients=[email_fornecedor]
        )
        
        empresa = dados_proposta.get('dados', {}).get('razaoSocial', 'N/A')
        processo = dados_proposta.get('processo', 'N/A')
        
        msg.html = f"""
        <html>
        <body style="font-family: Arial, sans-serif;">
            <div style="background-color: #f0f0f0; padding: 20px; border-radius: 10px;">
                <h2 style="color: #333;">‚úÖ Proposta Recebida com Sucesso!</h2>
                
                <div style="background-color: white; padding: 20px; border-radius: 5px; margin-top: 20px;">
                    <p>Prezado(a) {empresa},</p>
                    
                    <p>Confirmamos o recebimento de sua proposta para o processo <strong>{processo}</strong>.</p>
                    
                    <div style="background-color: #e3f2fd; padding: 15px; border-radius: 5px; margin: 20px 0;">
                        <p style="margin: 0;"><strong>Protocolo:</strong> {protocolo}</p>
                        <p style="margin: 0;"><strong>Data/Hora:</strong> {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}</p>
                    </div>
                    
                    <p>Sua proposta ser√° analisada e voc√™ ser√° notificado sobre o andamento do processo.</p>
                    
                    <p>Guarde este protocolo para futuras consultas.</p>
                    
                    <p style="margin-top: 30px;">Atenciosamente,<br>
                    Equipe de Suprimentos</p>
                </div>
                
                <p style="text-align: center; color: #666; margin-top: 20px;">
                    <small>Este √© um e-mail autom√°tico, por favor n√£o responda.</small>
                </p>
            </div>
        </body>
        </html>
        """
        
        mail.send(msg)
        logger.info(f"Email de confirma√ß√£o enviado para {email_fornecedor}")
        return True
        
    except Exception as e:
        logger.error(f"Erro ao enviar confirma√ß√£o ao fornecedor: {str(e)}")
        return False

def enviar_notificacao_novo_processo(processo):
    """Envia notifica√ß√£o aos fornecedores sobre novo processo"""
    try:
        # Buscar fornecedores cadastrados
        fornecedores = list(fornecedores_db.values())
        
        if not fornecedores:
            logger.info("Nenhum fornecedor cadastrado para notificar")
            return
        
        for fornecedor in fornecedores:
            if fornecedor.get('email'):
                msg = Message(
                    subject=f"Novo Processo de Licita√ß√£o - {processo['numero']}",
                    recipients=[fornecedor['email']]
                )
                
                prazo = datetime.fromisoformat(processo['prazo'].replace('Z', '+00:00'))
                prazo_formatado = prazo.strftime('%d/%m/%Y √†s %H:%M')
                
                msg.html = f"""
                <html>
                <body style="font-family: Arial, sans-serif;">
                    <div style="background-color: #f0f0f0; padding: 20px; border-radius: 10px;">
                        <h2 style="color: #333;">üì¢ Novo Processo de Licita√ß√£o</h2>
                        
                        <div style="background-color: white; padding: 20px; border-radius: 5px; margin-top: 20px;">
                            <p>Prezado(a) {fornecedor.get('razaoSocial', 'Fornecedor')},</p>
                            
                            <p>Informamos que foi aberto um novo processo de licita√ß√£o que pode ser de seu interesse:</p>
                            
                            <div style="background-color: #e3f2fd; padding: 15px; border-radius: 5px; margin: 20px 0;">
                                <p><strong>N√∫mero:</strong> {processo['numero']}</p>
                                <p><strong>Objeto:</strong> {processo['objeto']}</p>
                                <p><strong>Modalidade:</strong> {processo['modalidade']}</p>
                                <p><strong>Prazo para envio:</strong> {prazo_formatado}</p>
                            </div>
                            
                            <p>Para participar, acesse o portal de propostas atrav√©s do link:</p>
                            
                            <div style="text-align: center; margin: 30px 0;">
                                <a href="{request.url_root}portal-propostas?processo={processo['numero']}" 
                                   style="background-color: #4CAF50; color: white; padding: 15px 30px; 
                                          text-decoration: none; border-radius: 5px; display: inline-block;">
                                    Acessar Portal de Propostas
                                </a>
                            </div>
                            
                            <p>N√£o perca esta oportunidade!</p>
                            
                            <p style="margin-top: 30px;">Atenciosamente,<br>
                            Equipe de Suprimentos</p>
                        </div>
                    </div>
                </body>
                </html>
                """
                
                mail.send(msg)
                logger.info(f"Notifica√ß√£o enviada para {fornecedor['email']}")
                
    except Exception as e:
        logger.error(f"Erro ao enviar notifica√ß√µes: {str(e)}")

def verificar_prazos_proximos():
    """Verifica processos com prazo pr√≥ximo e envia lembretes"""
    try:
        agora = datetime.now()
        tres_dias = timedelta(days=3)
        
        for processo in processos_db.values():
            prazo = datetime.fromisoformat(processo['prazo'].replace('Z', '+00:00'))
            
            # Se faltam 3 dias ou menos
            if prazo > agora and (prazo - agora) <= tres_dias:
                # Buscar fornecedores que n√£o enviaram proposta
                propostas_processo = [p for p in propostas_db.values() if p['processo'] == processo['numero']]
                cnpjs_com_proposta = [p['dados']['cnpj'] for p in propostas_processo if 'dados' in p and 'cnpj' in p['dados']]
                
                for fornecedor in fornecedores_db.values():
                    if fornecedor.get('cnpj') not in cnpjs_com_proposta and fornecedor.get('email'):
                        enviar_lembrete_prazo(fornecedor, processo)
                        
    except Exception as e:
        logger.error(f"Erro ao verificar prazos: {str(e)}")

def enviar_lembrete_prazo(fornecedor, processo):
    """Envia lembrete de prazo para fornecedor"""
    try:
        msg = Message(
            subject=f"‚è∞ Lembrete: Prazo pr√≥ximo - Processo {processo['numero']}",
            recipients=[fornecedor['email']]
        )
        
        prazo = datetime.fromisoformat(processo['prazo'].replace('Z', '+00:00'))
        prazo_formatado = prazo.strftime('%d/%m/%Y √†s %H:%M')
        
        msg.html = f"""
        <html>
        <body style="font-family: Arial, sans-serif;">
            <div style="background-color: #fff3cd; padding: 20px; border-radius: 10px; border: 1px solid #ffeaa7;">
                <h2 style="color: #856404;">‚è∞ Lembrete de Prazo</h2>
                
                <div style="background-color: white; padding: 20px; border-radius: 5px; margin-top: 20px;">
                    <p>Prezado(a) {fornecedor.get('razaoSocial', 'Fornecedor')},</p>
                    
                    <p>Este √© um lembrete de que o prazo para envio de propostas para o processo abaixo est√° pr√≥ximo:</p>
                    
                    <div style="background-color: #ffeaa7; padding: 15px; border-radius: 5px; margin: 20px 0;">
                        <p><strong>Processo:</strong> {processo['numero']}</p>
                        <p><strong>Objeto:</strong> {processo['objeto']}</p>
                        <p><strong>PRAZO FINAL:</strong> {prazo_formatado}</p>
                    </div>
                    
                    <p><strong>Ainda n√£o recebemos sua proposta!</strong></p>
                    
                    <div style="text-align: center; margin: 30px 0;">
                        <a href="{request.url_root}portal-propostas?processo={processo['numero']}" 
                           style="background-color: #ff9800; color: white; padding: 15px 30px; 
                                  text-decoration: none; border-radius: 5px; display: inline-block;">
                            Enviar Proposta Agora
                        </a>
                    </div>
                    
                    <p style="color: #dc3545;"><strong>N√£o perca o prazo!</strong></p>
                </div>
            </div>
        </body>
        </html>
        """
        
        mail.send(msg)
        
        # Enviar c√≥pia para o comprador
        if EMAIL_CONFIG['destinatario_principal']:
            msg_comprador = Message(
                subject=f"Lembrete enviado - Processo {processo['numero']}",
                recipients=[EMAIL_CONFIG['destinatario_principal']]
            )
            msg_comprador.body = f"Lembrete de prazo enviado para {fornecedor['razaoSocial']} ({fornecedor['email']})"
            mail.send(msg_comprador)
            
        logger.info(f"Lembrete enviado para {fornecedor['email']}")
        
    except Exception as e:
        logger.error(f"Erro ao enviar lembrete: {str(e)}")

def salvar_proposta_arquivo(proposta_id, proposta_data):
    """Salva a proposta em arquivo JSON"""
    try:
        safe_id = str(proposta_id).replace('/', '_').replace('\\', '_').replace(':', '_')
        filename = os.path.join(PROPOSTAS_DIR, f'proposta_{safe_id}.json')
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(proposta_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Proposta salva: {filename}")
        return True
    except Exception as e:
        logger.error(f"Erro ao salvar proposta: {str(e)}")
        return False

def carregar_dados():
    """Carrega propostas, processos e fornecedores do diret√≥rio"""
    try:
        if not os.path.exists(PROPOSTAS_DIR):
            logger.info(f"Diret√≥rio {PROPOSTAS_DIR} n√£o existe ainda")
            return
            
        # Carrega propostas
        for filename in os.listdir(PROPOSTAS_DIR):
            if filename.endswith('.json') and filename.startswith('proposta_'):
                filepath = os.path.join(PROPOSTAS_DIR, filename)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        proposta = json.load(f)
                        propostas_db[proposta.get('protocolo', proposta.get('id'))] = proposta
                except Exception as e:
                    logger.error(f"Erro ao carregar {filename}: {e}")
        
        logger.info(f"Carregadas {len(propostas_db)} propostas")
        
        # Carrega processos se existir arquivo
        processos_file = os.path.join(PROPOSTAS_DIR, 'processos.json')
        if os.path.exists(processos_file):
            with open(processos_file, 'r', encoding='utf-8') as f:
                processos = json.load(f)
                for processo in processos:
                    processos_db[processo['numero']] = processo
                    
        # Carrega fornecedores se existir arquivo
        fornecedores_file = os.path.join(PROPOSTAS_DIR, 'fornecedores.json')
        if os.path.exists(fornecedores_file):
            with open(fornecedores_file, 'r', encoding='utf-8') as f:
                fornecedores = json.load(f)
                for fornecedor in fornecedores:
                    fornecedores_db[fornecedor.get('cnpj', fornecedor.get('id'))] = fornecedor
                    
    except Exception as e:
        logger.error(f"Erro ao carregar dados: {str(e)}")

# Carrega dados existentes ao iniciar
carregar_dados()

@app.route('/')
def home():
    """Serve a p√°gina principal ou informa√ß√µes da API"""
    static_index = os.path.join(STATIC_DIR, 'index.html')
    if os.path.exists(static_index):
        return send_from_directory(STATIC_DIR, 'index.html')
    else:
        return jsonify({
            "message": "Sistema de Propostas - API v2.0",
            "status": "online",
            "endpoints": [
                "/portal-propostas",
                "/api/enviar-proposta",
                "/api/status",
                "/api/propostas/listar",
                "/api/processos/listar",
                "/api/processos/<numero>",
                "/api/criar-processo",
                "/api/cadastrar-fornecedor"
            ],
            "email_configurado": bool(app.config['MAIL_USERNAME']),
            "total_propostas": len(propostas_db),
            "total_processos": len(processos_db),
            "total_fornecedores": len(fornecedores_db)
        }), 200

@app.route('/portal-propostas')
def portal_propostas():
    """Serve o portal de propostas"""
    portal_path = os.path.join(STATIC_DIR, 'portal-propostas-novo.html')
    if os.path.exists(portal_path):
        return send_from_directory(STATIC_DIR, 'portal-propostas-novo.html')
    else:
        return render_template_string("""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Portal de Propostas</title>
            <meta charset="UTF-8">
        </head>
        <body>
            <h1>Portal de Propostas</h1>
            <p>O arquivo portal-propostas-novo.html n√£o foi encontrado.</p>
            <p>Por favor, fa√ßa o upload do arquivo para a pasta static/</p>
            <hr>
            <p><a href="/api/status">Verificar Status da API</a></p>
        </body>
        </html>
        """)

@app.route('/api/status', methods=['GET'])
def api_status():
    """Verifica o status do servidor"""
    return jsonify({
        "status": "online",
        "timestamp": datetime.now().isoformat(),
        "propostas_total": len(propostas_db),
        "processos_total": len(processos_db),
        "fornecedores_total": len(fornecedores_db),
        "encoding": "UTF-8",
        "versao": "2.0.0",
        "email_configurado": bool(app.config['MAIL_USERNAME']),
        "diretorios": {
            "propostas": os.path.exists(PROPOSTAS_DIR),
            "static": os.path.exists(STATIC_DIR)
        }
    }), 200

@app.route('/api/enviar-proposta', methods=['POST'])
def enviar_proposta():
    """Recebe proposta do portal novo"""
    try:
        data = request.get_json(force=True)
        
        # Gera protocolo √∫nico
        protocolo = data.get('protocolo')
        if not protocolo:
            protocolo = f"PROP-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{str(uuid.uuid4())[:8].upper()}"
        
        # Verifica se fornecedor j√° enviou proposta para este processo
        processo_numero = data.get('processo', 'N/A')
        cnpj = data.get('dados', {}).get('cnpj', '')
        
        # Verifica propostas existentes
        for proposta in propostas_db.values():
            if (proposta.get('processo') == processo_numero and 
                proposta.get('dados', {}).get('cnpj') == cnpj):
                return jsonify({
                    "success": False,
                    "erro": "Sua empresa j√° enviou uma proposta para este processo"
                }), 400
        
        # Estrutura completa da proposta
        proposta = {
            "protocolo": protocolo,
            "data_envio": datetime.now().isoformat(),
            "processo": processo_numero,
            "status": "recebida",
            "dados": data.get('dados', {}),
            "resumo": data.get('resumo', {}),
            "tecnica": data.get('tecnica', {}),
            "comercial": data.get('comercial', {})
        }
        
        # Armazena na mem√≥ria
        propostas_db[protocolo] = proposta
        
        # Salva em arquivo
        if salvar_proposta_arquivo(protocolo, proposta):
            logger.info(f"Nova proposta recebida: {protocolo}")
            
            # Envia email com anexos
            email_enviado = enviar_email_proposta(protocolo, proposta)
            
            return jsonify({
                "success": True,
                "protocolo": protocolo,
                "mensagem": "Proposta enviada com sucesso!",
                "email_enviado": email_enviado,
                "data": datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
                "proposta_resumo": {
                    "protocolo": protocolo,
                    "processo": processo_numero,
                    "empresa": proposta['dados'].get('razaoSocial', 'N/A'),
                    "cnpj": cnpj,
                    "valor": f"R$ {proposta['comercial'].get('valorTotal', '0,00')}",
                    "data": proposta['data_envio'],
                    "dados": proposta
                }
            }), 201
        else:
            return jsonify({
                "success": False,
                "erro": "Erro ao salvar proposta"
            }), 500
            
    except Exception as e:
        logger.error(f"Erro ao processar proposta: {str(e)}")
        return jsonify({
            "success": False,
            "erro": "Erro ao calcular estat√≠sticas"
        }, 500

# Endpoint alternativo com query parameter para evitar problemas de codifica√ß√£o
@app.route('/api/fornecedor/estatisticas', methods=['GET'])
def estatisticas_fornecedor_query():
    """Retorna estat√≠sticas espec√≠ficas do fornecedor via query parameter"""
    cnpj = request.args.get('cnpj', '')
    if not cnpj:
        return jsonify({"success": False, "erro": "CNPJ n√£o fornecido"}), 400
    return estatisticas_fornecedor(cnpj)

# ===== FIM DOS NOVOS ENDPOINTS =====

@app.route('/api/criar-processo', methods=['POST'])
def criar_processo():
    """Cria um novo processo e notifica fornecedores"""
    try:
        data = request.get_json(force=True)
        
        processo = {
            "id": str(uuid.uuid4()),
            "numero": data.get('numero', ''),
            "objeto": data.get('objeto', ''),
            "modalidade": data.get('modalidade', ''),
            "prazo": data.get('prazo', ''),
            "dataCadastro": datetime.now().isoformat(),
            "criadoPor": data.get('criadoPor', ''),
            "notificarFornecedores": data.get('notificarFornecedores', False)
        }
        
        # Salva processo
        processos_db[processo['numero']] = processo
        
        # Salvar em arquivo
        processos_file = os.path.join(PROPOSTAS_DIR, 'processos.json')
        with open(processos_file, 'w', encoding='utf-8') as f:
            json.dump(list(processos_db.values()), f, ensure_ascii=False, indent=2)
        
        # Notificar fornecedores se solicitado
        if processo['notificarFornecedores']:
            enviar_notificacao_novo_processo(processo)
        
        return jsonify({
            "success": True,
            "processo": processo,
            "mensagem": "Processo criado com sucesso!"
        }), 201
        
    except Exception as e:
        logger.error(f"Erro ao criar processo: {str(e)}")
        return jsonify({
            "success": False,
            "erro": "Erro ao criar processo"
        }), 500

@app.route('/api/cadastrar-fornecedor', methods=['POST'])
def cadastrar_fornecedor():
    """Cadastra um novo fornecedor"""
    try:
        data = request.get_json(force=True)
        
        cnpj = data.get('cnpj', '')
        
        # Verifica se j√° existe
        if cnpj in fornecedores_db:
            return jsonify({
                "success": False,
                "erro": "Fornecedor j√° cadastrado"
            }), 400
        
        fornecedor = {
            "id": str(uuid.uuid4()),
            "cnpj": cnpj,
            "razaoSocial": data.get('razaoSocial', ''),
            "email": data.get('email', ''),
            "telefone": data.get('telefone', ''),
            "endereco": data.get('endereco', ''),
            "cidade": data.get('cidade', ''),
            "estado": data.get('estado', ''),
            "cep": data.get('cep', ''),
            "responsavel": data.get('responsavel', ''),
            "dataCadastro": datetime.now().isoformat(),
            "ativo": True
        }
        
        # Salva fornecedor
        fornecedores_db[cnpj] = fornecedor
        
        # Salvar em arquivo
        fornecedores_file = os.path.join(PROPOSTAS_DIR, 'fornecedores.json')
        with open(fornecedores_file, 'w', encoding='utf-8') as f:
            json.dump(list(fornecedores_db.values()), f, ensure_ascii=False, indent=2)
        
        return jsonify({
            "success": True,
            "fornecedor": fornecedor,
            "mensagem": "Fornecedor cadastrado com sucesso!"
        }), 201
        
    except Exception as e:
        logger.error(f"Erro ao cadastrar fornecedor: {str(e)}")
        return jsonify({
            "success": False,
            "erro": "Erro ao cadastrar fornecedor"
        }), 500

# Endpoints para arquivos est√°ticos movidos para o final

@app.errorhandler(404)
def nao_encontrado(e):
    return jsonify({"erro": "Endpoint n√£o encontrado"}), 404

@app.errorhandler(500)
def erro_interno(e):
    return jsonify({"erro": "Erro interno do servidor"}), 500

@app.route('/api/download/proposta/<protocolo>/<tipo>', methods=['GET'])
def download_proposta(protocolo, tipo):
    """Download de proposta em Excel ou Word usando as mesmas fun√ß√µes dos anexos de e-mail"""
    try:
        # Buscar proposta
        proposta = None
        for p in propostas_db.values():
            if p.get('protocolo') == protocolo:
                proposta = p
                break
        
        if not proposta:
            return jsonify({
                "success": False,
                "erro": "Proposta n√£o encontrada"
            }), 404
        
        if tipo == 'excel':
            # Gerar Excel usando a mesma fun√ß√£o do e-mail
            excel_buffer = gerar_excel_proposta(proposta)
            
            # Configurar response
            response = send_file(
                excel_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f'proposta_{protocolo}_completa.xlsx'
            )
            return response
            
        elif tipo == 'word':
            # Gerar Word usando a mesma fun√ß√£o do e-mail
            word_buffer = gerar_word_proposta(proposta)
            
            # Configurar response
            response = send_file(
                word_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'proposta_{protocolo}_tecnica.docx'
            )
            return response
            
        else:
            return jsonify({
                "success": False,
                "erro": "Tipo inv√°lido. Use 'excel' ou 'word'"
            }), 400
            
    except Exception as e:
        logger.error(f"Erro ao fazer download: {str(e)}")
        return jsonify({
            "success": False,
            "erro": f"Erro ao gerar arquivo: {str(e)}"
        }), 500

# ===== EXTENS√ÉO DO BACKEND - NOVOS ENDPOINTS =====
# Este arquivo cont√©m apenas os novos endpoints para os m√≥dulos TR
# Para integrar: copiar e colar no FINAL do backend_propostas.py

# ===== NOVOS ENDPOINTS PARA M√ìDULOS ADICIONAIS =====

@app.route('/api/trs', methods=['GET'])
def listar_trs():
    """Listar Termos de Refer√™ncia"""
    try:
        # Carregar TRs de arquivo JSON
        trs_file = 'data/trs.json'
        if os.path.exists(trs_file):
            with open(trs_file, 'r', encoding='utf-8') as f:
                trs = json.load(f)
        else:
            trs = []
        
        # Filtrar por usu√°rio se especificado
        user_id = request.args.get('user_id')
        if user_id:
            trs = [tr for tr in trs if tr.get('criado_por') == user_id]
        
        return jsonify({'success': True, 'trs': trs, 'total': len(trs)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/trs', methods=['POST'])
def criar_tr():
    """Criar novo Termo de Refer√™ncia"""
    try:
        data = request.get_json()
        
        # Carregar TRs existentes
        trs_file = 'data/trs.json'
        if os.path.exists(trs_file):
            with open(trs_file, 'r', encoding='utf-8') as f:
                trs = json.load(f)
        else:
            trs = []
        
        # Criar novo TR
        novo_tr = {
            'id': f"TR-{len(trs) + 1:03d}",
            'titulo': data.get('titulo', ''),
            'objetivo': data.get('objetivo', ''),
            'situacao_atual': data.get('situacao_atual', ''),
            'problemas': data.get('problemas', ''),
            'necessidades': data.get('necessidades', ''),
            'modalidade': data.get('modalidade', 'concorrencia'),
            'especificacoes': data.get('especificacoes', ''),
            'normas': data.get('normas', ''),
            'prazo_execucao': data.get('prazo_execucao', 0),
            'prazo_garantia': data.get('prazo_garantia', 0),
            'servicos': data.get('servicos', []),
            'criado_por': data.get('criado_por', 'sistema'),
            'criado_em': datetime.now().isoformat(),
            'atualizado_em': datetime.now().isoformat(),
            'status': data.get('status', 'rascunho')
        }
        
        trs.append(novo_tr)
        
        # Salvar TRs
        os.makedirs('data', exist_ok=True)
        with open(trs_file, 'w', encoding='utf-8') as f:
            json.dump(trs, f, ensure_ascii=False, indent=2)
        
        return jsonify({'success': True, 'tr': novo_tr})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/trs/<tr_id>', methods=['GET'])
def obter_tr(tr_id):
    """Obter TR espec√≠fico"""
    try:
        trs_file = 'data/trs.json'
        if os.path.exists(trs_file):
            with open(trs_file, 'r', encoding='utf-8') as f:
                trs = json.load(f)
            
            tr = next((t for t in trs if t['id'] == tr_id), None)
            if tr:
                return jsonify({'success': True, 'tr': tr})
            else:
                return jsonify({'success': False, 'error': 'TR n√£o encontrado'})
        else:
            return jsonify({'success': False, 'error': 'Nenhum TR encontrado'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/trs/<tr_id>', methods=['PUT'])
def atualizar_tr(tr_id):
    """Atualizar TR existente"""
    try:
        data = request.get_json()
        
        trs_file = 'data/trs.json'
        if os.path.exists(trs_file):
            with open(trs_file, 'r', encoding='utf-8') as f:
                trs = json.load(f)
        else:
            return jsonify({'success': False, 'error': 'TR n√£o encontrado'})
        
        # Encontrar e atualizar TR
        for i, tr in enumerate(trs):
            if tr['id'] == tr_id:
                # Manter dados originais
                trs[i].update({
                    'titulo': data.get('titulo', tr['titulo']),
                    'objetivo': data.get('objetivo', tr['objetivo']),
                    'situacao_atual': data.get('situacao_atual', tr['situacao_atual']),
                    'problemas': data.get('problemas', tr.get('problemas', '')),
                    'necessidades': data.get('necessidades', tr.get('necessidades', '')),
                    'modalidade': data.get('modalidade', tr.get('modalidade', 'concorrencia')),
                    'especificacoes': data.get('especificacoes', tr.get('especificacoes', '')),
                    'normas': data.get('normas', tr.get('normas', '')),
                    'prazo_execucao': data.get('prazo_execucao', tr.get('prazo_execucao', 0)),
                    'prazo_garantia': data.get('prazo_garantia', tr.get('prazo_garantia', 0)),
                    'servicos': data.get('servicos', tr.get('servicos', [])),
                    'status': data.get('status', tr['status']),
                    'atualizado_em': datetime.now().isoformat()
                })
                
                # Salvar altera√ß√µes
                with open(trs_file, 'w', encoding='utf-8') as f:
                    json.dump(trs, f, ensure_ascii=False, indent=2)
                
                return jsonify({'success': True, 'tr': trs[i]})
        
        return jsonify({'success': False, 'error': 'TR n√£o encontrado'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/trs/<tr_id>', methods=['DELETE'])
def excluir_tr(tr_id):
    """Excluir TR"""
    try:
        trs_file = 'data/trs.json'
        if os.path.exists(trs_file):
            with open(trs_file, 'r', encoding='utf-8') as f:
                trs = json.load(f)
        else:
            return jsonify({'success': False, 'error': 'TR n√£o encontrado'})
        
        # Encontrar e remover TR
        trs_filtrados = [tr for tr in trs if tr['id'] != tr_id]
        
        if len(trs_filtrados) < len(trs):
            # TR foi removido
            with open(trs_file, 'w', encoding='utf-8') as f:
                json.dump(trs_filtrados, f, ensure_ascii=False, indent=2)
            
            return jsonify({'success': True, 'message': 'TR exclu√≠do com sucesso'})
        else:
            return jsonify({'success': False, 'error': 'TR n√£o encontrado'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/dashboard/requisitante/<user_id>')
def dashboard_requisitante(user_id):
    """Dashboard espec√≠fico para requisitante"""
    try:
        # Carregar dados dos TRs
        trs_file = 'data/trs.json'
        if os.path.exists(trs_file):
            with open(trs_file, 'r', encoding='utf-8') as f:
                trs = json.load(f)
        else:
            trs = []
        
        # Filtrar TRs do usu√°rio
        meus_trs = [tr for tr in trs if tr.get('criado_por') == user_id]
        trs_aprovados = [tr for tr in meus_trs if tr.get('status') == 'aprovado']
        
        # Simular dados de propostas em an√°lise
        propostas_analise = 2
        pareceres_pendentes = 1
        
        dados = {
            'meus_trs': len(meus_trs),
            'trs_aprovados': len(trs_aprovados),
            'propostas_analise': propostas_analise,
            'pareceres_pendentes': pareceres_pendentes,
            'trs_recentes': meus_trs[-5:] if meus_trs else []
        }
        
        return jsonify({'success': True, 'dados': dados})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/dashboard/comprador/<user_id>')
def dashboard_comprador(user_id):
    """Dashboard espec√≠fico para comprador"""
    try:
        # Carregar dados dos TRs
        trs_file = 'data/trs.json'
        if os.path.exists(trs_file):
            with open(trs_file, 'r', encoding='utf-8') as f:
                trs = json.load(f)
        else:
            trs = []
        
        # Simular dados do comprador
        trs_pendentes = len([tr for tr in trs if tr.get('status') == 'analise'])
        concorrencias_ativas = 2
        propostas_recebidas = 8
        valor_total = 1250000
        
        dados = {
            'trs_pendentes': trs_pendentes,
            'concorrencias_ativas': concorrencias_ativas,
            'propostas_recebidas': propostas_recebidas,
            'valor_total': valor_total,
            'trs_para_analise': [tr for tr in trs if tr.get('status') in ['rascunho', 'analise']]
        }
        
        return jsonify({'success': True, 'dados': dados})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/propostas/tecnicas', methods=['GET'])
def listar_propostas_tecnicas():
    """Listar propostas t√©cnicas (sem valores comerciais)"""
    try:
        # Simular dados de propostas t√©cnicas
        propostas_tecnicas = [
            {
                'id': 'PROP-001',
                'fornecedor': 'Construtora ABC Ltda',
                'tr_id': 'TR-001',
                'metodologia': 'Metodologia construtiva com t√©cnicas modernas...',
                'cronograma': '120 dias',
                'equipe_tecnica': 'Eng. Jo√£o Silva (CREA 12345) + equipe de 15 profissionais',
                'experiencia': '10 obras similares nos √∫ltimos 5 anos',
                'certificacoes': 'ISO 9001, ISO 14001',
                'recebida_em': datetime.now().isoformat(),
                'status': 'analise_tecnica'
            },
            {
                'id': 'PROP-002',
                'fornecedor': 'Engenharia XYZ S.A.',
                'tr_id': 'TR-001',
                'metodologia': 'Abordagem sustent√°vel com materiais eco-friendly...',
                'cronograma': '90 dias',
                'equipe_tecnica': 'Eng. Maria Santos (CREA 67890) + equipe especializada',
                'experiencia': '15 obras similares, incluindo projetos governamentais',
                'certificacoes': 'ISO 9001, PBQP-H',
                'recebida_em': datetime.now().isoformat(),
                'status': 'analise_tecnica'
            }
        ]
        
        return jsonify({'success': True, 'propostas': propostas_tecnicas})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/propostas/<proposta_id>/parecer', methods=['POST'])
def emitir_parecer_tecnico(proposta_id):
    """Emitir parecer t√©cnico sobre proposta"""
    try:
        data = request.get_json()
        
        parecer = {
            'id': f"PARECER-{proposta_id}",
            'proposta_id': proposta_id,
            'avaliacao_metodologia': data.get('avaliacao_metodologia', ''),
            'avaliacao_cronograma': data.get('avaliacao_cronograma', ''),
            'avaliacao_equipe': data.get('avaliacao_equipe', ''),
            'pontos_positivos': data.get('pontos_positivos', ''),
            'pontos_negativos': data.get('pontos_negativos', ''),
            'recomendacao': data.get('recomendacao', ''),  # aprovado/rejeitado/condicional
            'observacoes': data.get('observacoes', ''),
            'emitido_por': data.get('emitido_por', 'requisitante'),
            'emitido_em': datetime.now().isoformat()
        }
        
        # Salvar parecer
        pareceres_file = 'data/pareceres.json'
        if os.path.exists(pareceres_file):
            with open(pareceres_file, 'r', encoding='utf-8') as f:
                pareceres = json.load(f)
        else:
            pareceres = []
        
        pareceres.append(parecer)
        
        os.makedirs('data', exist_ok=True)
        with open(pareceres_file, 'w', encoding='utf-8') as f:
            json.dump(pareceres, f, ensure_ascii=False, indent=2)
        
        return jsonify({'success': True, 'parecer': parecer})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/pareceres', methods=['GET'])
def listar_pareceres():
    """Listar pareceres t√©cnicos"""
    try:
        pareceres_file = 'data/pareceres.json'
        if os.path.exists(pareceres_file):
            with open(pareceres_file, 'r', encoding='utf-8') as f:
                pareceres = json.load(f)
        else:
            pareceres = []
        
        # Filtrar por usu√°rio se especificado
        user_id = request.args.get('user_id')
        if user_id:
            pareceres = [p for p in pareceres if p.get('emitido_por') == user_id]
        
        return jsonify({'success': True, 'pareceres': pareceres})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/concorrencias', methods=['GET'])
def listar_concorrencias():
    """Listar concorr√™ncias ativas"""
    try:
        # Simular dados de concorr√™ncias
        concorrencias = [
            {
                'id': 'CONC-001',
                'numero': '001/2025',
                'objeto': 'Reforma do Centro de Sa√∫de Municipal',
                'modalidade': 'Concorr√™ncia',
                'valor_estimado': 850000.00,
                'prazo_proposta': '2025-02-15',
                'dias_restantes': 15,
                'status': 'aberta',
                'tr_id': 'TR-001'
            },
            {
                'id': 'CONC-002',
                'numero': '002/2025',
                'objeto': 'Constru√ß√£o de Quadra Poliesportiva',
                'modalidade': 'Tomada de Pre√ßos',
                'valor_estimado': 450000.00,
                'prazo_proposta': '2025-02-20',
                'dias_restantes': 20,
                'status': 'aberta',
                'tr_id': 'TR-002'
            },
            {
                'id': 'CONC-003',
                'numero': '003/2025',
                'objeto': 'Pavimenta√ß√£o de Ruas do Bairro Centro',
                'modalidade': 'Concorr√™ncia',
                'valor_estimado': 1200000.00,
                'prazo_proposta': '2025-02-25',
                'dias_restantes': 25,
                'status': 'aberta',
                'tr_id': 'TR-003'
            }
        ]
        
        return jsonify({'success': True, 'concorrencias': concorrencias})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/inicializar-dados', methods=['POST'])
def inicializar_dados_exemplo_api():
    """Inicializar dados de exemplo para demonstra√ß√£o"""
    try:
        # Criar diret√≥rio de dados
        os.makedirs('data', exist_ok=True)
        
        # TRs de exemplo
        trs_exemplo = [
            {
                'id': 'TR-001',
                'titulo': 'Reforma do Centro de Sa√∫de Municipal',
                'objetivo': 'Reforma completa das instala√ß√µes do Centro de Sa√∫de para melhorar o atendimento √† popula√ß√£o',
                'situacao_atual': 'Instala√ß√µes antigas com necessidade de moderniza√ß√£o',
                'modalidade': 'concorrencia',
                'servicos': [
                    {'item': 1, 'descricao': 'Demoli√ß√£o de paredes internas', 'unidade': 'm¬≤', 'quantidade': 150},
                    {'item': 2, 'descricao': 'Constru√ß√£o de novas paredes', 'unidade': 'm¬≤', 'quantidade': 200},
                    {'item': 3, 'descricao': 'Instala√ß√£o el√©trica completa', 'unidade': 'vb', 'quantidade': 1}
                ],
                'criado_por': 'requisitante1',
                'criado_em': datetime.now().isoformat(),
                'status': 'aprovado'
            },
            {
                'id': 'TR-002',
                'titulo': 'Constru√ß√£o de Quadra Poliesportiva',
                'objetivo': 'Constru√ß√£o de quadra poliesportiva coberta para atividades esportivas da comunidade',
                'situacao_atual': 'Terreno dispon√≠vel, necessidade de espa√ßo esportivo',
                'modalidade': 'tomada_precos',
                'servicos': [
                    {'item': 1, 'descricao': 'Terraplanagem e funda√ß√£o', 'unidade': 'm¬≤', 'quantidade': 800},
                    {'item': 2, 'descricao': 'Estrutura met√°lica da cobertura', 'unidade': 'm¬≤', 'quantidade': 600},
                    {'item': 3, 'descricao': 'Piso esportivo', 'unidade': 'm¬≤', 'quantidade': 600}
                ],
                'criado_por': 'requisitante1',
                'criado_em': datetime.now().isoformat(),
                'status': 'analise'
            }
        ]
        
        # Salvar TRs
        with open('data/trs.json', 'w', encoding='utf-8') as f:
            json.dump(trs_exemplo, f, ensure_ascii=False, indent=2)
        
        # Pareceres de exemplo
        pareceres_exemplo = []
        with open('data/pareceres.json', 'w', encoding='utf-8') as f:
            json.dump(pareceres_exemplo, f, ensure_ascii=False, indent=2)
        
        return jsonify({'success': True, 'message': 'Dados de exemplo inicializados'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

# ===== FIM DOS NOVOS ENDPOINTS =====

# ===== ENDPOINT CATCH-ALL (DEVE SER O √öLTIMO) =====
@app.route('/<path:filename>')
def serve_static(filename):
    """Serve arquivos est√°ticos com tratamento de erro"""
    try:
        # Previne path traversal
        if '..' in filename or filename.startswith('/'):
            return jsonify({"erro": "Caminho inv√°lido"}), 400
            
        file_path = os.path.join(STATIC_DIR, filename)
        if os.path.exists(file_path) and os.path.isfile(file_path):
            return send_from_directory(STATIC_DIR, filename)
        else:
            return jsonify({"erro": "Arquivo n√£o encontrado"}), 404
    except Exception as e:
        logger.error(f"Erro ao servir arquivo {filename}: {e}")
        return jsonify({"erro": "Erro ao acessar arquivo"}), 500

if __name__ == '__main__':
    logger.info("Iniciando servidor de propostas v2.0.0...")
    logger.info(f"Diret√≥rio de trabalho: {os.getcwd()}")
    logger.info(f"Email configurado: {bool(app.config['MAIL_USERNAME'])}")
    logger.info(f"Destinat√°rio principal: {EMAIL_CONFIG['destinatario_principal']}")
    
    # Verificar prazos periodicamente (em produ√ß√£o usar scheduler apropriado)
    # verificar_prazos_proximos()
    
    port = int(os.environ.get('PORT', 5000))
    debug_mode = os.environ.get('DEBUG', 'False').lower() == 'true'
    
    app.run(
        host='0.0.0.0',
        port=port,
        debug=debug_mode
    ) False,
            "erro": "Erro interno ao processar proposta",
            "detalhes": str(e)
        }, 500

@app.route('/api/propostas/listar', methods=['GET'])
def listar_propostas():
    """Lista todas as propostas com formato compat√≠vel com o frontend"""
    try:
        # Filtros opcionais
        processo = request.args.get('processo')
        cnpj = request.args.get('cnpj')
        
        propostas_lista = []
        
        for proposta in propostas_db.values():
            # Aplicar filtros se especificados
            if processo and proposta.get('processo') != processo:
                continue
            if cnpj and proposta.get('dados', {}).get('cnpj') != cnpj:
                continue
            
            # Formatar proposta para o frontend
            proposta_formatada = {
                "protocolo": proposta.get('protocolo', ''),
                "processo": proposta.get('processo', ''),
                "empresa": proposta.get('dados', {}).get('razaoSocial', 'N/A'),
                "cnpj": proposta.get('dados', {}).get('cnpj', ''),
                "valor": f"R$ {proposta.get('comercial', {}).get('valorTotal', '0,00')}",
                "data": proposta.get('data_envio', ''),
                "dados": proposta
            }
            
            propostas_lista.append(proposta_formatada)
        
        # Ordena por data
        propostas_lista.sort(key=lambda x: x.get('data', ''), reverse=True)
        
        return jsonify({
            "propostas": propostas_lista,
            "total": len(propostas_lista)
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao listar propostas: {str(e)}")
        return jsonify({
            "erro": "Erro ao listar propostas"
        }), 500

@app.route('/api/processos/<numero>', methods=['GET'])
def obter_processo(numero):
    """Obt√©m informa√ß√µes de um processo espec√≠fico"""
    if numero in processos_db:
        return jsonify(processos_db[numero]), 200
    else:
        # Retorna dados padr√£o se n√£o encontrar
        return jsonify({
            "numero": numero,
            "objeto": "Processo n√£o cadastrado",
            "modalidade": "N√£o informada",
            "prazo": datetime.now().isoformat()
        }), 200

@app.route('/api/processos/listar', methods=['GET'])
def listar_processos():
    """Lista todos os processos"""
    return jsonify({
        "processos": list(processos_db.values()),
        "total": len(processos_db)
    }), 200

# ===== NOVOS ENDPOINTS PARA √ÅREA DO FORNECEDOR =====

@app.route('/api/processos/ativos', methods=['GET'])
def listar_processos_ativos():
    """Lista apenas processos com prazo n√£o vencido"""
    try:
        agora = datetime.now()
        processos_ativos = []
        
        for processo in processos_db.values():
            try:
                # Converter prazo para datetime
                prazo = datetime.fromisoformat(processo['prazo'].replace('Z', '+00:00'))
                
                # Verificar se ainda est√° ativo
                if prazo > agora:
                    # Calcular dias restantes
                    dias_restantes = (prazo - agora).days
                    
                    processo_ativo = {
                        "numero": processo['numero'],
                        "objeto": processo['objeto'],
                        "modalidade": processo['modalidade'],
                        "prazo": processo['prazo'],
                        "prazo_formatado": prazo.strftime('%d/%m/%Y %H:%M'),
                        "dias_restantes": dias_restantes,
                        "status": "ativo"
                    }
                    processos_ativos.append(processo_ativo)
                    
            except Exception as e:
                logger.error(f"Erro ao processar processo {processo.get('numero', 'N/A')}: {e}")
                continue
        
        # Ordenar por prazo (mais pr√≥ximos primeiro)
        processos_ativos.sort(key=lambda x: x['prazo'])
        
        return jsonify({
            "success": True,
            "processos": processos_ativos,
            "total": len(processos_ativos)
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao listar processos ativos: {str(e)}")
        return jsonify({
            "success": False,
            "erro": "Erro ao listar processos ativos"
        }), 500

@app.route('/api/propostas/fornecedor/<cnpj>', methods=['GET'])
def listar_propostas_fornecedor(cnpj):
    """Lista propostas de um fornecedor espec√≠fico"""
    try:
        # Remover formata√ß√£o do CNPJ para compara√ß√£o
        cnpj_limpo = cnpj.replace('.', '').replace('/', '').replace('-', '')
        
        propostas_fornecedor = []
        
        for proposta in propostas_db.values():
            proposta_cnpj = proposta.get('dados', {}).get('cnpj', '')
            proposta_cnpj_limpo = proposta_cnpj.replace('.', '').replace('/', '').replace('-', '')
            
            if proposta_cnpj_limpo == cnpj_limpo:
                # Buscar informa√ß√µes do processo
                processo = processos_db.get(proposta.get('processo', ''), {})
                
                proposta_formatada = {
                    "protocolo": proposta.get('protocolo', ''),
                    "processo": proposta.get('processo', ''),
                    "processo_objeto": processo.get('objeto', 'N/A'),
                    "data_envio": proposta.get('data_envio', ''),
                    "data_envio_formatada": datetime.fromisoformat(proposta.get('data_envio', '')).strftime('%d/%m/%Y %H:%M') if proposta.get('data_envio') else '',
                    "valor_total": proposta.get('comercial', {}).get('valorTotal', '0,00'),
                    "status": proposta.get('status', 'enviada'),
                    "empresa": proposta.get('dados', {}).get('razaoSocial', ''),
                    "processo_prazo": processo.get('prazo', ''),
                    "processo_encerrado": datetime.fromisoformat(processo.get('prazo', '').replace('Z', '+00:00')) < datetime.now() if processo.get('prazo') else False
                }
                
                propostas_fornecedor.append(proposta_formatada)
        
        # Ordenar por data de envio (mais recentes primeiro)
        propostas_fornecedor.sort(key=lambda x: x.get('data_envio', ''), reverse=True)
        
        return jsonify({
            "success": True,
            "propostas": propostas_fornecedor,
            "total": len(propostas_fornecedor),
            "cnpj": cnpj
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao listar propostas do fornecedor {cnpj}: {str(e)}")
        return jsonify({
            "success": False,
            "erro": "Erro ao listar propostas do fornecedor"
        }), 500

# Endpoint alternativo com query parameter para evitar problemas de codifica√ß√£o
@app.route('/api/propostas/fornecedor', methods=['GET'])
def listar_propostas_fornecedor_query():
    """Lista propostas de um fornecedor espec√≠fico via query parameter"""
    cnpj = request.args.get('cnpj', '')
    if not cnpj:
        return jsonify({"success": False, "erro": "CNPJ n√£o fornecido"}), 400
    return listar_propostas_fornecedor(cnpj)

@app.route('/api/fornecedor/estatisticas/<cnpj>', methods=['GET'])
def estatisticas_fornecedor(cnpj):
    """Retorna estat√≠sticas espec√≠ficas do fornecedor"""
    try:
        # Remover formata√ß√£o do CNPJ
        cnpj_limpo = cnpj.replace('.', '').replace('/', '').replace('-', '')
        
        # Contar propostas do fornecedor
        propostas_fornecedor = []
        for proposta in propostas_db.values():
            proposta_cnpj = proposta.get('dados', {}).get('cnpj', '')
            proposta_cnpj_limpo = proposta_cnpj.replace('.', '').replace('/', '').replace('-', '')
            
            if proposta_cnpj_limpo == cnpj_limpo:
                propostas_fornecedor.append(proposta)
        
        # Contar processos ativos
        agora = datetime.now()
        processos_ativos = 0
        for processo in processos_db.values():
            try:
                prazo = datetime.fromisoformat(processo['prazo'].replace('Z', '+00:00'))
                if prazo > agora:
                    processos_ativos += 1
            except:
                continue
        
        # Contar prazos pr√≥ximos (pr√≥ximos 7 dias)
        uma_semana = agora + timedelta(days=7)
        prazos_proximos = 0
        for processo in processos_db.values():
            try:
                prazo = datetime.fromisoformat(processo['prazo'].replace('Z', '+00:00'))
                if agora < prazo <= uma_semana:
                    prazos_proximos += 1
            except:
                continue
        
        # Calcular valor total das propostas
        valor_total = 0
        for proposta in propostas_fornecedor:
            valor_str = proposta.get('comercial', {}).get('valorTotal', '0,00')
            try:
                # Remover formata√ß√£o e converter para float
                valor_limpo = valor_str.replace('R$', '').replace('.', '').replace(',', '.').strip()
                valor_total += float(valor_limpo)
            except:
                continue
        
        return jsonify({
            "success": True,
            "estatisticas": {
                "processos_ativos": processos_ativos,
                "propostas_enviadas": len(propostas_fornecedor),
                "prazos_proximos": prazos_proximos,
                "valor_total": f"{valor_total:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao calcular estat√≠sticas do fornecedor {cnpj}: {str(e)}")
        return jsonify({
            "success":
