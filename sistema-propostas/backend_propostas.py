#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from flask import Flask, request, jsonify, send_from_directory, render_template_string
from flask_cors import CORS
from datetime import datetime
import json
import os
import uuid
import logging
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from io import BytesIO
import base64

# Configuração do Flask
app = Flask(__name__)
app.config['JSON_AS_ASCII'] = False
app.config['JSONIFY_MIMETYPE'] = 'application/json; charset=utf-8'
CORS(app)

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Diretórios
PROPOSTAS_DIR = 'propostas'
STATIC_DIR = 'static'

# Criar diretórios se não existirem
for dir_name in [PROPOSTAS_DIR, STATIC_DIR]:
    try:
        if not os.path.exists(dir_name):
            os.makedirs(dir_name)
            logger.info(f"Diretório {dir_name} criado")
    except Exception as e:
        logger.error(f"Erro ao criar diretório {dir_name}: {e}")

# Base de dados em memória
propostas_db = {}
processos_db = {}

# Configurações de email (usar variáveis de ambiente em produção)
EMAIL_CONFIG = {
    'server': os.environ.get('EMAIL_SERVER', 'smtp.gmail.com'),
    'port': int(os.environ.get('EMAIL_PORT', 587)),
    'user': os.environ.get('EMAIL_USER', ''),
    'password': os.environ.get('EMAIL_PASS', ''),
    'destinatario': os.environ.get('EMAIL_SUPRIMENTOS', '')
}

def gerar_proposta_excel(protocolo, dados_proposta):
    """Gera proposta comercial em Excel"""
    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Proposta Comercial"
        
        # Estilos
        header_font = Font(bold=True, size=14, color="FFFFFF")
        header_fill = PatternFill(start_color="667EEA", end_color="667EEA", fill_type="solid")
        sub_header_font = Font(bold=True, size=12)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Cabeçalho
        ws.merge_cells('A1:F1')
        ws['A1'] = f"PROPOSTA COMERCIAL - {protocolo}"
        ws['A1'].font = Font(bold=True, size=16)
        ws['A1'].alignment = Alignment(horizontal='center')
        
        # Dados da empresa
        ws['A3'] = "DADOS DA EMPRESA"
        ws['A3'].font = header_font
        ws['A3'].fill = header_fill
        ws.merge_cells('A3:F3')
        
        dados_empresa = dados_proposta.get('dados', {}).get('dados', {})
        ws['A4'] = "Razão Social:"
        ws['B4'] = dados_empresa.get('razaoSocial', '')
        ws['A5'] = "CNPJ:"
        ws['B5'] = dados_empresa.get('cnpj', '')
        ws['A6'] = "E-mail:"
        ws['B6'] = dados_empresa.get('email', '')
        ws['D4'] = "Telefone:"
        ws['E4'] = dados_empresa.get('telefone', '')
        ws['D5'] = "Responsável:"
        ws['E5'] = dados_empresa.get('respTecnico', '')
        
        # Resumo da proposta
        ws['A8'] = "RESUMO DA PROPOSTA"
        ws['A8'].font = header_font
        ws['A8'].fill = header_fill
        ws.merge_cells('A8:F8')
        
        comercial = dados_proposta.get('comercial', {})
        ws['A9'] = "Valor Total:"
        ws['B9'] = f"R$ {comercial.get('valorTotal', '0,00')}"
        ws['A10'] = "Prazo de Execução:"
        ws['B10'] = dados_proposta.get('tecnica', {}).get('prazoExecucao', '')
        ws['A11'] = "Validade da Proposta:"
        ws['B11'] = comercial.get('validadeProposta', '60 dias')
        
        # Serviços
        linha = 13
        if comercial.get('servicos'):
            ws[f'A{linha}'] = "PLANILHA DE SERVIÇOS"
            ws[f'A{linha}'].font = header_font
            ws[f'A{linha}'].fill = header_fill
            ws.merge_cells(f'A{linha}:F{linha}')
            linha += 1
            
            # Cabeçalho da tabela
            headers = ['Item', 'Descrição', 'Unid', 'Qtd', 'Preço Unit', 'Total']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=linha, column=col, value=header)
                cell.font = sub_header_font
                cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
                cell.border = border
            linha += 1
            
            # Dados dos serviços
            for servico in comercial['servicos']:
                for col, valor in enumerate(servico[:6], 1):
                    cell = ws.cell(row=linha, column=col, value=valor)
                    cell.border = border
                linha += 1
            
            # Total dos serviços
            ws[f'E{linha}'] = "TOTAL:"
            ws[f'F{linha}'] = f"R$ {comercial.get('totalServicos', '0,00')}"
            ws[f'F{linha}'].font = Font(bold=True)
            linha += 2
        
        # Mão de obra
        if comercial.get('maoObra'):
            ws[f'A{linha}'] = "MÃO DE OBRA"
            ws[f'A{linha}'].font = header_font
            ws[f'A{linha}'].fill = header_fill
            ws.merge_cells(f'A{linha}:G{linha}')
            linha += 1
            
            headers = ['Função', 'Qtd', 'Meses', 'Salário', 'Enc.(%)', 'Total Mensal', 'Total Geral']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=linha, column=col, value=header)
                cell.font = sub_header_font
                cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
                cell.border = border
            linha += 1
            
            for mo in comercial['maoObra']:
                for col, valor in enumerate(mo[:7], 1):
                    cell = ws.cell(row=linha, column=col, value=valor)
                    cell.border = border
                linha += 1
            
            ws[f'F{linha}'] = "TOTAL:"
            ws[f'G{linha}'] = f"R$ {comercial.get('totalMaoObra', '0,00')}"
            ws[f'G{linha}'].font = Font(bold=True)
            linha += 2
        
        # BDI
        if comercial.get('bdi'):
            ws[f'A{linha}'] = "COMPOSIÇÃO DO BDI"
            ws[f'A{linha}'].font = header_font
            ws[f'A{linha}'].fill = header_fill
            ws.merge_cells(f'A{linha}:C{linha}')
            linha += 1
            
            headers = ['Componente', 'Percentual (%)', 'Valor (R$)']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=linha, column=col, value=header)
                cell.font = sub_header_font
                cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
                cell.border = border
            linha += 1
            
            for item_bdi in comercial['bdi']:
                for col, valor in enumerate(item_bdi[:3], 1):
                    cell = ws.cell(row=linha, column=col, value=valor)
                    cell.border = border
                linha += 1
            
            ws[f'A{linha}'] = "BDI TOTAL:"
            ws[f'B{linha}'] = f"{comercial.get('bdiPercentual', '0')}%"
            ws[f'C{linha}'] = f"R$ {comercial.get('bdiValor', '0,00')}"
            for col in range(1, 4):
                ws.cell(row=linha, column=col).font = Font(bold=True)
            linha += 2
        
        # Valor final
        ws[f'A{linha}'] = "VALOR TOTAL DA PROPOSTA"
        ws[f'A{linha}'].font = Font(bold=True, size=14)
        ws.merge_cells(f'A{linha}:C{linha}')
        linha += 1
        ws[f'A{linha}'] = f"R$ {comercial.get('valorTotal', '0,00')}"
        ws[f'A{linha}'].font = Font(bold=True, size=16, color="667EEA")
        ws.merge_cells(f'A{linha}:C{linha}')
        
        # Ajustar largura das colunas
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Salvar em BytesIO
        excel_buffer = BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)
        
        return excel_buffer
        
    except Exception as e:
        logger.error(f"Erro ao gerar Excel: {str(e)}")
        return None

def gerar_proposta_word(protocolo, dados_proposta):
    """Gera proposta técnica em Word"""
    try:
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Título
        titulo = doc.add_heading(f'PROPOSTA TÉCNICA', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Protocolo
        p = doc.add_paragraph()
        p.add_run(f'Protocolo: {protocolo}').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        p = doc.add_paragraph()
        p.add_run(f'Data: {datetime.now().strftime("%d/%m/%Y")}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Dados da empresa
        doc.add_heading('1. DADOS DA EMPRESA', level=1)
        dados_empresa = dados_proposta.get('dados', {}).get('dados', {})
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid'
        
        cells = [
            ('Razão Social:', dados_empresa.get('razaoSocial', '')),
            ('CNPJ:', dados_empresa.get('cnpj', '')),
            ('Endereço:', dados_empresa.get('endereco', '')),
            ('Cidade/UF:', dados_empresa.get('cidade', '')),
            ('Responsável Técnico:', dados_empresa.get('respTecnico', ''))
        ]
        
        for i, (label, value) in enumerate(cells):
            table.cell(i, 0).text = label
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            table.cell(i, 1).text = value
        
        doc.add_paragraph()
        
        # Objeto
        doc.add_heading('2. OBJETO DA CONCORRÊNCIA', level=1)
        tecnica = dados_proposta.get('tecnica', {})
        doc.add_paragraph(tecnica.get('objetoConcorrencia', ''))
        
        # Escopo
        doc.add_heading('3. ESCOPO DOS SERVIÇOS', level=1)
        if tecnica.get('escopoInclusos'):
            doc.add_heading('3.1 Serviços Inclusos', level=2)
            doc.add_paragraph(tecnica.get('escopoInclusos', ''))
        
        if tecnica.get('escopoExclusos'):
            doc.add_heading('3.2 Serviços Exclusos', level=2)
            doc.add_paragraph(tecnica.get('escopoExclusos', ''))
        
        # Metodologia
        doc.add_heading('4. METODOLOGIA DE EXECUÇÃO', level=1)
        doc.add_paragraph(tecnica.get('metodologia', ''))
        
        if tecnica.get('sequenciaExecucao'):
            doc.add_heading('4.1 Sequência de Execução', level=2)
            doc.add_paragraph(tecnica.get('sequenciaExecucao', ''))
        
        # Prazo
        doc.add_heading('5. PRAZO DE EXECUÇÃO', level=1)
        doc.add_paragraph(f"Prazo total: {tecnica.get('prazoExecucao', '')}")
        
        if tecnica.get('prazoMobilizacao'):
            doc.add_paragraph(f"Prazo de mobilização: {tecnica.get('prazoMobilizacao', '')}")
        
        # Cronograma
        if tecnica.get('cronograma'):
            doc.add_heading('6. CRONOGRAMA DE EXECUÇÃO', level=1)
            
            table = doc.add_table(rows=len(tecnica['cronograma']) + 1, cols=4)
            table.style = 'Light Grid'
            
            # Cabeçalho
            headers = ['Atividade', 'Duração', 'Início', 'Fim']
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # Dados
            for i, atividade in enumerate(tecnica['cronograma'], 1):
                for j, valor in enumerate(atividade[:4]):
                    table.cell(i, j).text = str(valor)
        
        # Equipe
        if tecnica.get('equipe'):
            doc.add_heading('7. EQUIPE TÉCNICA', level=1)
            
            table = doc.add_table(rows=len(tecnica['equipe']) + 1, cols=3)
            table.style = 'Light Grid'
            
            headers = ['Função', 'Quantidade', 'Tempo (meses)']
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            for i, profissional in enumerate(tecnica['equipe'], 1):
                for j, valor in enumerate(profissional[:3]):
                    table.cell(i, j).text = str(valor)
        
        # Experiência
        if tecnica.get('experiencia'):
            doc.add_heading('8. EXPERIÊNCIA DA EMPRESA', level=1)
            
            for i, obra in enumerate(tecnica['experiencia'], 1):
                p = doc.add_paragraph()
                p.add_run(f'{i}. {obra[0]}').bold = True
                doc.add_paragraph(f'   Cliente: {obra[1]}')
                doc.add_paragraph(f'   Valor: {obra[2]}')
                doc.add_paragraph(f'   Ano: {obra[3]}')
                doc.add_paragraph()
        
        # Salvar em BytesIO
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        
        return word_buffer
        
    except Exception as e:
        logger.error(f"Erro ao gerar Word: {str(e)}")
        return None

def enviar_email_proposta(protocolo, dados_proposta):
    """Envia email com os dados da proposta e anexos Excel e Word"""
    if not EMAIL_CONFIG['user'] or not EMAIL_CONFIG['password']:
        logger.warning("Configurações de email não definidas")
        return False
    
    try:
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['user']
        msg['To'] = EMAIL_CONFIG['destinatario']
        msg['Subject'] = f"Nova Proposta - {protocolo}"
        
        # Extrai informações principais
        empresa = dados_proposta.get('dados', {}).get('razaoSocial', 'N/A')
        cnpj = dados_proposta.get('dados', {}).get('cnpj', 'N/A')
        valor_total = dados_proposta.get('comercial', {}).get('valorTotal', '0,00')
        processo = dados_proposta.get('processo', 'N/A')
        prazo_execucao = dados_proposta.get('tecnica', {}).get('prazoExecucao', 'N/A')
        
        corpo = f"""
        NOVA PROPOSTA RECEBIDA
        
        Protocolo: {protocolo}
        Processo: {processo}
        
        DADOS DA EMPRESA:
        Empresa: {empresa}
        CNPJ: {cnpj}
        Email: {dados_proposta.get('dados', {}).get('email', 'N/A')}
        Telefone: {dados_proposta.get('dados', {}).get('telefone', 'N/A')}
        
        RESUMO DA PROPOSTA:
        Valor Total: R$ {valor_total}
        Prazo de Execução: {prazo_execucao}
        
        Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
        
        ==========================================
        ANEXOS:
        1. Proposta Comercial (Excel)
        2. Proposta Técnica (Word)
        3. Dados Completos (JSON)
        ==========================================
        """
        
        msg.attach(MIMEText(corpo, 'plain'))
        
        # Gerar e anexar Excel
        excel_buffer = gerar_proposta_excel(protocolo, dados_proposta)
        if excel_buffer:
            excel_attachment = MIMEBase('application', 'vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            excel_attachment.set_payload(excel_buffer.read())
            encoders.encode_base64(excel_attachment)
            excel_attachment.add_header('Content-Disposition', f'attachment; filename=proposta_comercial_{protocolo}.xlsx')
            msg.attach(excel_attachment)
        
        # Gerar e anexar Word
        word_buffer = gerar_proposta_word(protocolo, dados_proposta)
        if word_buffer:
            word_attachment = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
            word_attachment.set_payload(word_buffer.read())
            encoders.encode_base64(word_attachment)
            word_attachment.add_header('Content-Disposition', f'attachment; filename=proposta_tecnica_{protocolo}.docx')
            msg.attach(word_attachment)
        
        # Anexar dados completos como JSON
        json_attachment = MIMEBase('application', 'json')
        json_attachment.set_payload(json.dumps(dados_proposta, ensure_ascii=False, indent=2).encode('utf-8'))
        encoders.encode_base64(json_attachment)
        json_attachment.add_header('Content-Disposition', f'attachment; filename=proposta_completa_{protocolo}.json')
        msg.attach(json_attachment)
        
        # Envia o email
        server = smtplib.SMTP(EMAIL_CONFIG['server'], EMAIL_CONFIG['port'])
        server.starttls()
        server.login(EMAIL_CONFIG['user'], EMAIL_CONFIG['password'])
        server.send_message(msg)
        server.quit()
        
        logger.info(f"Email enviado com sucesso para proposta {protocolo} com anexos Excel e Word")
        
        # Enviar email de confirmação para o fornecedor
        enviar_email_confirmacao_fornecedor(protocolo, dados_proposta)
        
        return True
        
    except Exception as e:
        logger.error(f"Erro ao enviar email: {str(e)}")
        return False

def enviar_email_confirmacao_fornecedor(protocolo, dados_proposta):
    """Envia email de confirmação para o fornecedor"""
    try:
        email_fornecedor = dados_proposta.get('dados', {}).get('email')
        if not email_fornecedor:
            return False
        
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['user']
        msg['To'] = email_fornecedor
        msg['Subject'] = f"Confirmação de Envio - Proposta {protocolo}"
        
        empresa = dados_proposta.get('dados', {}).get('razaoSocial', 'N/A')
        processo = dados_proposta.get('processo', 'N/A')
        valor_total = dados_proposta.get('comercial', {}).get('valorTotal', '0,00')
        
        corpo = f"""
        Prezado(a) {empresa},
        
        Confirmamos o recebimento de sua proposta!
        
        DADOS DO ENVIO:
        Protocolo: {protocolo}
        Processo: {processo}
        Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
        Valor Total: R$ {valor_total}
        
        IMPORTANTE:
        - Sua proposta foi registrada com sucesso
        - Guarde este protocolo para acompanhamento
        - Não é possível editar a proposta após o envio
        - Para revisões, entre em contato com o setor de compras
        
        Atenciosamente,
        Setor de Licitações
        """
        
        msg.attach(MIMEText(corpo, 'plain'))
        
        server = smtplib.SMTP(EMAIL_CONFIG['server'], EMAIL_CONFIG['port'])
        server.starttls()
        server.login(EMAIL_CONFIG['user'], EMAIL_CONFIG['password'])
        server.send_message(msg)
        server.quit()
        
        logger.info(f"Email de confirmação enviado para {email_fornecedor}")
        return True
        
    except Exception as e:
        logger.error(f"Erro ao enviar confirmação: {str(e)}")
        return False

def salvar_proposta_arquivo(proposta_id, proposta_data):
    """Salva a proposta em arquivo JSON"""
    try:
        # Garante que o ID seja válido para nome de arquivo
        safe_id = str(proposta_id).replace('/', '_').replace('\\', '_').replace(':', '_')
        filename = os.path.join(PROPOSTAS_DIR, f'proposta_{safe_id}.json')
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(proposta_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Proposta salva: {filename}")
        return True
    except Exception as e:
        logger.error(f"Erro ao salvar proposta: {str(e)}")
        return False

def carregar_dados():
    """Carrega propostas e processos do diretório"""
    try:
        if not os.path.exists(PROPOSTAS_DIR):
            logger.info(f"Diretório {PROPOSTAS_DIR} não existe ainda")
            return
            
        # Carrega propostas
        for filename in os.listdir(PROPOSTAS_DIR):
            if filename.endswith('.json') and filename.startswith('proposta_'):
                filepath = os.path.join(PROPOSTAS_DIR, filename)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        proposta = json.load(f)
                        propostas_db[proposta.get('protocolo', proposta.get('id'))] = proposta
                except Exception as e:
                    logger.error(f"Erro ao carregar {filename}: {e}")
        
        logger.info(f"Carregadas {len(propostas_db)} propostas")
        
        # Carrega processos se existir arquivo
        processos_file = os.path.join(PROPOSTAS_DIR, 'processos.json')
        if os.path.exists(processos_file):
            with open(processos_file, 'r', encoding='utf-8') as f:
                processos = json.load(f)
                for processo in processos:
                    processos_db[processo['numero']] = processo
                    
    except Exception as e:
        logger.error(f"Erro ao carregar dados: {str(e)}")

# Carrega dados existentes ao iniciar
carregar_dados()

@app.route('/')
def home():
    """Serve a página principal ou informações da API"""
    static_index = os.path.join(STATIC_DIR, 'index.html')
    if os.path.exists(static_index):
        return send_from_directory(STATIC_DIR, 'index.html')
    else:
        return jsonify({
            "message": "Sistema de Propostas - API",
            "status": "online",
            "endpoints": [
                "/portal-propostas",
                "/api/enviar-proposta",
                "/api/status",
                "/api/propostas/listar",
                "/api/processos/listar"
            ]
        }), 200

@app.route('/portal-propostas')
def portal_propostas():
    """Serve o portal de propostas"""
    portal_path = os.path.join(STATIC_DIR, 'portal-propostas-novo.html')
    if os.path.exists(portal_path):
        return send_from_directory(STATIC_DIR, 'portal-propostas-novo.html')
    else:
        # Se o arquivo não existir, retorna um HTML básico
        return render_template_string("""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Portal de Propostas</title>
            <meta charset="UTF-8">
        </head>
        <body>
            <h1>Portal de Propostas</h1>
            <p>O arquivo portal-propostas-novo.html não foi encontrado.</p>
            <p>Por favor, faça o upload do arquivo para a pasta static/</p>
            <hr>
            <p><a href="/api/status">Verificar Status da API</a></p>
        </body>
        </html>
        """)

@app.route('/api/status', methods=['GET'])
def api_status():
    """Verifica o status do servidor"""
    return jsonify({
        "status": "online",
        "timestamp": datetime.now().isoformat(),
        "propostas_total": len(propostas_db),
        "processos_total": len(processos_db),
        "encoding": "UTF-8",
        "versao": "2.0.1",
        "email_configurado": bool(EMAIL_CONFIG['user']),
        "diretorios": {
            "propostas": os.path.exists(PROPOSTAS_DIR),
            "static": os.path.exists(STATIC_DIR)
        }
    }), 200

@app.route('/api/enviar-proposta', methods=['POST'])
def enviar_proposta():
    """Recebe proposta do portal novo"""
    try:
        data = request.get_json(force=True)
        
        # Extrair CNPJ e processo
        cnpj = data.get('dados', {}).get('cnpj', '')
        processo = data.get('processo', '')
        
        # Verificar se fornecedor já enviou proposta para este processo
        for proposta_existente in propostas_db.values():
            if (proposta_existente.get('dados', {}).get('dados', {}).get('cnpj') == cnpj and 
                proposta_existente.get('processo') == processo and
                not proposta_existente.get('revisao', False)):
                return jsonify({
                    "success": False,
                    "erro": "Sua empresa já enviou uma proposta para este processo. Entre em contato com o comprador para solicitar autorização de revisão."
                }), 400
        
        # Gera protocolo único
        protocolo = data.get('protocolo')
        if not protocolo:
            protocolo = f"PROP-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{str(uuid.uuid4())[:8].upper()}"
        
        # Estrutura completa da proposta
        proposta = {
            "protocolo": protocolo,
            "data_envio": datetime.now().isoformat(),
            "processo": processo,
            "status": "recebida",
            "revisao": False,
            "numero_revisao": 0,
            "dados": data.get('dados', {}),
            "resumo": data.get('resumo', {}),
            "tecnica": data.get('tecnica', {}),
            "comercial": data.get('comercial', {})
        }
        
        # Armazena na memória
        propostas_db[protocolo] = proposta
        
        # Salva em arquivo
        if salvar_proposta_arquivo(protocolo, proposta):
            logger.info(f"Nova proposta recebida: {protocolo}")
            
            # Tenta enviar email (não bloqueia se falhar)
            email_enviado = enviar_email_proposta(protocolo, proposta)
            
            return jsonify({
                "success": True,
                "protocolo": protocolo,
                "mensagem": "Proposta enviada com sucesso!",
                "email_enviado": email_enviado,
                "data": datetime.now().strftime('%d/%m/%Y %H:%M:%S')
            }), 201
        else:
            return jsonify({
                "success": False,
                "erro": "Erro ao salvar proposta"
            }), 500
            
    except Exception as e:
        logger.error(f"Erro ao processar proposta: {str(e)}")
        return jsonify({
            "success": False,
            "erro": "Erro interno ao processar proposta",
            "detalhes": str(e)
        }), 500

@app.route('/api/propostas/listar', methods=['GET'])
def listar_propostas():
    """Lista todas as propostas"""
    try:
        # Filtros opcionais
        processo = request.args.get('processo')
        cnpj = request.args.get('cnpj')
        
        propostas_lista = list(propostas_db.values())
        
        # Aplica filtros
        if processo:
            propostas_lista = [p for p in propostas_lista if p.get('processo') == processo]
        if cnpj:
            propostas_lista = [p for p in propostas_lista if p.get('dados', {}).get('cnpj') == cnpj]
        
        # Ordena por data
        propostas_lista.sort(key=lambda x: x.get('data_envio', ''), reverse=True)
        
        return jsonify({
            "propostas": propostas_lista,
            "total": len(propostas_lista)
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao listar propostas: {str(e)}")
        return jsonify({
            "erro": "Erro ao listar propostas"
        }), 500

@app.route('/api/processos/<numero>', methods=['GET'])
def obter_processo(numero):
    """Obtém informações de um processo específico"""
    if numero in processos_db:
        return jsonify(processos_db[numero]), 200
    else:
        # Retorna dados padrão se não encontrar
        return jsonify({
            "numero": numero,
            "objeto": "Processo não cadastrado",
            "prazo": datetime.now().isoformat()
        }), 200

@app.route('/api/processos/listar', methods=['GET'])
def listar_processos():
    """Lista todos os processos"""
    return jsonify({
        "processos": list(processos_db.values()),
        "total": len(processos_db)
    }), 200

@app.route('/api/propostas/autorizar-revisao', methods=['POST'])
def autorizar_revisao():
    """Autoriza fornecedor a enviar revisão de proposta"""
    try:
        data = request.get_json(force=True)
        cnpj = data.get('cnpj')
        processo = data.get('processo')
        autorizado_por = data.get('autorizado_por', 'admin')
        
        if not cnpj or not processo:
            return jsonify({
                "success": False,
                "erro": "CNPJ e processo são obrigatórios"
            }), 400
        
        # Encontrar proposta original
        proposta_original = None
        for prop in propostas_db.values():
            if (prop.get('dados', {}).get('dados', {}).get('cnpj') == cnpj and 
                prop.get('processo') == processo and
                not prop.get('revisao', False)):
                proposta_original = prop
                break
        
        if not proposta_original:
            return jsonify({
                "success": False,
                "erro": "Proposta original não encontrada"
            }), 404
        
        # Criar autorização de revisão
        autorizacao = {
            "id": str(uuid.uuid4()),
            "cnpj": cnpj,
            "processo": processo,
            "autorizado_por": autorizado_por,
            "data_autorizacao": datetime.now().isoformat(),
            "utilizada": False
        }
        
        # Salvar autorização (em produção, usar banco de dados)
        autorizacoes = json.loads(os.environ.get('AUTORIZACOES_REVISAO', '[]'))
        autorizacoes.append(autorizacao)
        os.environ['AUTORIZACOES_REVISAO'] = json.dumps(autorizacoes)
        
        logger.info(f"Autorização de revisão criada para CNPJ {cnpj} no processo {processo}")
        
        return jsonify({
            "success": True,
            "autorizacao_id": autorizacao["id"],
            "mensagem": "Fornecedor autorizado a enviar revisão"
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao autorizar revisão: {str(e)}")
        return jsonify({
            "success": False,
            "erro": "Erro ao processar autorização"
        }), 500

@app.route('/api/processos/notificar-fornecedores', methods=['POST'])
def notificar_fornecedores_novo_processo():
    """Notifica fornecedores cadastrados sobre novo processo"""
    try:
        data = request.get_json(force=True)
        processo = data.get('processo')
        
        if not processo:
            return jsonify({
                "success": False,
                "erro": "Dados do processo são obrigatórios"
            }), 400
        
        # Buscar fornecedores cadastrados (em produção, buscar do banco)
        # Por enquanto, simular com lista
        fornecedores_emails = data.get('fornecedores', [])
        
        if not fornecedores_emails:
            # Buscar todos os fornecedores do sistema
            # Em produção, isso viria do banco de dados
            return jsonify({
                "success": True,
                "mensagem": "Nenhum fornecedor para notificar"
            }), 200
        
        # Enviar email para cada fornecedor
        enviados = 0
        falhas = 0
        
        for email_fornecedor in fornecedores_emails:
            try:
                msg = MIMEMultipart()
                msg['From'] = EMAIL_CONFIG['user']
                msg['To'] = email_fornecedor
                msg['Subject'] = f"Novo Processo de Licitação - {processo.get('numero', '')}"
                
                # URL do portal
                portal_url = f"https://seu-dominio.onrender.com/portal-propostas?processo={processo.get('numero', '')}"
                
                corpo = f"""
                Prezado Fornecedor,
                
                Informamos que foi aberto um novo processo de licitação que pode ser do seu interesse.
                
                DADOS DO PROCESSO:
                Número: {processo.get('numero', '')}
                Objeto: {processo.get('objeto', '')}
                Modalidade: {processo.get('modalidade', '')}
                Prazo para envio de propostas: {processo.get('prazo', '')}
                
                Para participar deste processo, acesse o portal de propostas:
                {portal_url}
                
                IMPORTANTE:
                - Cada empresa pode enviar apenas uma proposta por processo
                - Após o envio, não será possível editar a proposta
                - Mantenha seus dados cadastrais atualizados
                
                Em caso de dúvidas, entre em contato com o setor de licitações.
                
                Atenciosamente,
                Setor de Licitações
                """
                
                msg.attach(MIMEText(corpo, 'plain'))
                
                server = smtplib.SMTP(EMAIL_CONFIG['server'], EMAIL_CONFIG['port'])
                server.starttls()
                server.login(EMAIL_CONFIG['user'], EMAIL_CONFIG['password'])
                server.send_message(msg)
                server.quit()
                
                enviados += 1
                logger.info(f"Notificação enviada para {email_fornecedor}")
                
            except Exception as e:
                falhas += 1
                logger.error(f"Erro ao notificar {email_fornecedor}: {str(e)}")
        
        return jsonify({
            "success": True,
            "enviados": enviados,
            "falhas": falhas,
            "total": len(fornecedores_emails)
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao notificar fornecedores: {str(e)}")
        return jsonify({
            "success": False,
            "erro": "Erro ao processar notificações"
        }), 500

@app.route('/api/processos/criar', methods=['POST'])
def criar_processo():
    """Cria novo processo e notifica fornecedores"""
    try:
        data = request.get_json(force=True)
        
        # Criar processo
        novo_processo = {
            "id": str(uuid.uuid4()),
            "numero": data.get('numero'),
            "objeto": data.get('objeto'),
            "modalidade": data.get('modalidade'),
            "prazo": data.get('prazo'),
            "dataCadastro": datetime.now().isoformat(),
            "criadoPor": data.get('criadoPor', 'admin')
        }
        
        # Salvar processo
        processos_db[novo_processo['numero']] = novo_processo
        
        # Salvar em arquivo
        processos_file = os.path.join(PROPOSTAS_DIR, 'processos.json')
        processos_list = list(processos_db.values())
        with open(processos_file, 'w', encoding='utf-8') as f:
            json.dump(processos_list, f, ensure_ascii=False, indent=2)
        
        # Buscar emails de fornecedores ativos
        fornecedores_emails = []
        # Em produção, buscar do banco de dados
        # Por enquanto, usar lista vazia ou receber do frontend
        
        # Notificar fornecedores
        if fornecedores_emails:
            notificar_fornecedores_novo_processo()
        
        return jsonify({
            "success": True,
            "processo": novo_processo,
            "mensagem": "Processo criado com sucesso!"
        }), 201
        
    except Exception as e:
        logger.error(f"Erro ao criar processo: {str(e)}")
        return jsonify({
            "success": False,
            "erro": "Erro ao criar processo"
        }), 500

# Endpoints para arquivos estáticos
@app.route('/<path:filename>')
def serve_static(filename):
    """Serve arquivos estáticos com tratamento de erro"""
    try:
        # Previne path traversal
        if '..' in filename or filename.startswith('/'):
            return jsonify({"erro": "Caminho inválido"}), 400
            
        file_path = os.path.join(STATIC_DIR, filename)
        if os.path.exists(file_path) and os.path.isfile(file_path):
            return send_from_directory(STATIC_DIR, filename)
        else:
            return jsonify({"erro": "Arquivo não encontrado"}), 404
    except Exception as e:
        logger.error(f"Erro ao servir arquivo {filename}: {e}")
        return jsonify({"erro": "Erro ao acessar arquivo"}), 500

@app.errorhandler(404)
def nao_encontrado(e):
    return jsonify({"erro": "Endpoint não encontrado"}), 404

@app.errorhandler(500)
def erro_interno(e):
    return jsonify({"erro": "Erro interno do servidor"}), 500

if __name__ == '__main__':
    logger.info("Iniciando servidor de propostas v2.0.1...")
    logger.info(f"Diretório de trabalho: {os.getcwd()}")
    logger.info(f"Arquivos no diretório: {os.listdir('.')}")
    
    port = int(os.environ.get('PORT', 5000))
    debug_mode = os.environ.get('DEBUG', 'False').lower() == 'true'
    
    app.run(
        host='0.0.0.0',
        port=port,
        debug=debug_mode
    )